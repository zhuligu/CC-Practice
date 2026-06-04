# -*- coding: utf-8 -*-
"""把 lesson HTML 文件转换为 Word (.docx)。

策略：
- 文本/列表/表格/代码块 → python-docx 原生构造
- 流程图（.flow-diagram）→ playwright 截图后嵌入图片
- 自动跳过：主题切换按钮、面包屑、lesson-badge 圆环

用法：
    python assets/html_to_docx.py lesson-01.html
    生成 lesson-01.docx
"""
import sys
from pathlib import Path
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from playwright.sync_api import sync_playwright


# ===== 截图 =====
def screenshot_flow_diagrams(html_path: Path, out_dir: Path) -> dict:
    """渲染 HTML，给每个 .flow-diagram 单独截图，返回 {idx: path}"""
    out_dir.mkdir(parents=True, exist_ok=True)
    images = {}
    with sync_playwright() as p:
        browser = p.chromium.launch()
        page = browser.new_page(viewport={"width": 1100, "height": 800}, device_scale_factor=2)
        page.goto(f"file:///{html_path.resolve().as_posix()}")
        page.wait_for_load_state("networkidle")
        # 用 light 主题截图（Word 看着舒服）
        page.evaluate("document.documentElement.setAttribute('data-theme','light')")
        page.wait_for_timeout(150)
        diagrams = page.query_selector_all(".flow-diagram")
        for i, d in enumerate(diagrams):
            img_path = out_dir / f"flow-{i:02d}.png"
            d.screenshot(path=str(img_path), omit_background=False)
            images[i] = img_path
            print(f"  screenshot: flow-{i:02d}.png  ({img_path.stat().st_size // 1024} KB)")
        browser.close()
    return images


# ===== 工具：给 run / paragraph 加底色 =====
def _shade(elem, fill_hex: str):
    pr = elem.get_or_add_rPr() if elem.tag.endswith("}r") else elem.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pr.append(shd)


def add_shaded_paragraph(doc, text, font="Consolas", size=9, fill="F4F4F4"):
    p = doc.add_paragraph()
    _shade(p._p, fill)
    # 边框
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "4")
        b.set(qn("w:color"), "DDDDDD")
        pbdr.append(b)
    p_pr.append(pbdr)
    # 文本（保留换行）
    lines = text.split("\n")
    for i, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = font
        run.font.size = Pt(size)
        if i < len(lines) - 1:
            run.add_break()
    return p


# ===== 内联运行（处理 <strong>/<em>/<code>/<a>/<br> 等） =====
def add_inline_runs(p, node):
    for child in node.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text.strip() or text == " ":
                p.add_run(text)
            continue
        tag = child.name
        if tag == "code":
            run = p.add_run(child.get_text())
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            _shade(run._element, "F4F4F4")
        elif tag in ("strong", "b"):
            run = p.add_run(child.get_text())
            run.bold = True
        elif tag in ("em", "i"):
            run = p.add_run(child.get_text())
            run.italic = True
        elif tag == "a":
            run = p.add_run(child.get_text())
            run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
            run.underline = True
        elif tag == "br":
            p.add_run().add_break()
        elif tag == "span":
            classes = child.get("class", []) or []
            if "tag" in classes:
                run = p.add_run(f"[{child.get_text()}]  ")
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
            elif "lesson-badge" in classes:
                # 跳过圆环徽章
                continue
            elif "icon" in classes:
                continue  # 主题按钮的小图标
            else:
                add_inline_runs(p, child)
        else:
            # 其他标签：递归取文本
            add_inline_runs(p, child)


# ===== 表格 =====
def process_table(doc, table):
    # 计算行/列
    rows = table.find_all("tr")
    if not rows:
        return
    n_cols = max(len(r.find_all(["th", "td"])) for r in rows)
    n_rows = len(rows)
    t = doc.add_table(rows=n_rows, cols=n_cols)
    t.style = "Light Grid Accent 1"
    for ri, tr in enumerate(rows):
        cells = tr.find_all(["th", "td"])
        for ci, cell in enumerate(cells):
            if ci >= n_cols:
                continue
            tc = t.rows[ri].cells[ci]
            # 清空默认段落
            tc.text = ""
            p = tc.paragraphs[0]
            add_inline_runs(p, cell)
            if cell.name == "th":
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph()  # 表后空行


# ===== Callout（提示框）=====
def process_callout(doc, node):
    classes = node.get("class", []) or []
    if "warn" in classes:
        fill = "FFF1F0"
        accent = RGBColor(0xFF, 0x4D, 0x4F)
    elif "ok" in classes:
        fill = "F6FFED"
        accent = RGBColor(0x52, 0xC1, 0x1A)
    else:
        fill = "FFFBE6"
        accent = RGBColor(0xFA, 0xDB, 0x14)
    # 标题段（粗体彩色）
    strong = node.find("strong")
    if strong:
        p = doc.add_paragraph()
        _shade(p._p, fill)
        run = p.add_run(strong.get_text())
        run.bold = True
        run.font.color.rgb = accent
        strong.extract()  # 移除避免重复
    # 内容段
    for child in node.children:
        if isinstance(child, NavigableString):
            continue
        if child.name == "p":
            p = doc.add_paragraph()
            _shade(p._p, fill)
            add_inline_runs(p, child)
    doc.add_paragraph()


# ===== 块级元素分发 =====
def process_block(doc, node, flow_images, flow_state):
    if isinstance(node, NavigableString):
        return
    tag = node.name
    if not tag:
        return
    classes = node.get("class", []) or []

    # 跳过：面包屑、主题按钮
    if "breadcrumb" in classes:
        return
    if tag == "button":
        return

    if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
        level = int(tag[1])
        # h1 里的 lesson-badge 已经被忽略；取剩余文本
        for badge in node.find_all("span", class_="lesson-badge"):
            badge.extract()
        text = node.get_text(strip=True)
        h = doc.add_heading(level=min(level, 4))
        run = h.add_run(text)
    elif tag == "p":
        if not node.get_text(strip=True) and not node.find("br"):
            return
        p = doc.add_paragraph()
        add_inline_runs(p, node)
    elif tag == "ul":
        for li in node.find_all("li", recursive=False):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, li)
    elif tag == "ol":
        for li in node.find_all("li", recursive=False):
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, li)
    elif tag == "pre":
        code = node.find("code")
        text = code.get_text() if code else node.get_text()
        add_shaded_paragraph(doc, text)
    elif tag == "blockquote":
        for child in node.children:
            if isinstance(child, NavigableString):
                continue
            if child.name == "p":
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                add_inline_runs(p, child)
                for run in p.runs:
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    elif tag == "table":
        process_table(doc, node)
    elif tag == "hr":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("— — —")
    elif tag == "div":
        if "callout" in classes:
            process_callout(doc, node)
        elif "flow-diagram" in classes:
            idx = flow_state["idx"]
            img_path = flow_images.get(idx)
            if img_path and img_path.exists():
                doc.add_picture(str(img_path), width=Cm(15))
                last_p = doc.paragraphs[-1]
                last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            flow_state["idx"] = idx + 1
        else:
            for child in node.children:
                process_block(doc, child, flow_images, flow_state)
    elif tag in ("main", "section", "article"):
        for child in node.children:
            process_block(doc, child, flow_images, flow_state)
    # 其他标签（script/style/button/...）忽略


# ===== 主流程 =====
def convert(html_path: Path):
    print(f"=== 转换 {html_path.name} ===")
    out_dir = Path("assets/_flow_images") / html_path.stem
    images = screenshot_flow_diagrams(html_path, out_dir)
    print(f"  共截图 {len(images)} 张流程图")

    soup = BeautifulSoup(html_path.read_text(encoding="utf-8"), "lxml")
    doc = Document()

    # 页边距
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    # 默认字体（中英文混排）
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")

    main = soup.find("main")
    if not main:
        main = soup.body or soup
    flow_state = {"idx": 0}
    for child in main.children:
        process_block(doc, child, images, flow_state)

    out_path = html_path.parent / f"{html_path.stem}.docx"
    doc.save(str(out_path))
    print(f"  ✓ saved: {out_path}  ({out_path.stat().st_size // 1024} KB)")
    return out_path


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("usage: python assets/html_to_docx.py <lesson-XX.html> [...]")
        sys.exit(1)
    for arg in sys.argv[1:]:
        convert(Path(arg))
